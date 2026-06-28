from pathlib import Path
from docx import Document as DocxDocument
from pypdf import PdfReader

from knowledge_engine.models import Document


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}

    def load_directories(self, directories: list[str | Path]) -> list[Document]:
        all_documents: list[Document] = []

        for directory in directories:
            all_documents.extend(self.load_directory(directory))

        return all_documents

    def load_directory(self, directory: str | Path) -> list[Document]:
        root = Path(directory)

        if not root.exists():
            raise FileNotFoundError(f"Directory not found: {root}")

        documents: list[Document] = []

        for path in root.rglob("*"):
            if not path.is_file():
                continue

            if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                continue

            text = self.load_text(path)

            if not text.strip():
                continue

            relative_path = str(path.relative_to(root))

            documents.append(
                Document(
                    text=text,
                    source=relative_path,
                    path=str(path),
                    document_type=self.infer_document_type(path, root),
                    metadata={
                        "extension": path.suffix.lower(),
                        "file_name": path.name,
                        "folder": str(path.parent.relative_to(root)),
                        "root": str(root),
                        "relative_path": relative_path,
                        "knowledge_source": root.name,
                    },
                )
            )

        return documents

    def load_text(self, path: str | Path) -> str:
        path = Path(path)
        ext = path.suffix.lower()

        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")

        if ext == ".docx":
            return self._read_docx(path)

        if ext == ".pdf":
            return self._read_pdf(path)

        return ""

    def _read_docx(self, path: Path) -> str:
        doc = DocxDocument(str(path))
        return "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

    def _read_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def infer_document_type(self, path: Path, root: Path) -> str:
        relative = str(path.relative_to(root)).lower().replace("\\", "/")

        if relative.startswith("agents/"):
            return "AgentDefinition"

        if relative.startswith("workflows/"):
            return "Workflow"

        if relative.startswith("rules/"):
            return "Rule"

        if relative.startswith("memory/"):
            return "Memory"

        if relative.startswith("aiducation/"):
            return "AIducation"

        if relative.startswith("architecture/adr/"):
            return "ADR"

        if "architecture" in relative:
            return "Architecture"

        if "system-specifications" in relative:
            return "SystemSpecification"

        if "business" in relative:
            return "Business"

        if "research" in relative:
            return "Research"

        if "threat" in relative or "intelligence" in relative:
            return "ThreatIntelligence"

        return "General"